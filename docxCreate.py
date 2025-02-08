import docx 
from json import loads
from llmConnect import generateDocx

def generateSampleDocx():
    doc = docx.Document()
    doc.add_heading('Document Title', 0)
    doc_para = doc.add_paragraph('This is a paragraph.')
    doc_para.add_run('This is a new run.')
    doc.save('test.docx')



def generateDocxFromDict(docxDict,file_name="testDataFromLLM.docx"):
    doc = docx.Document()
    for topic_number, topic_details in docxDict.items():
        doc.add_heading(topic_details["heading"])
        doc_para = doc.add_paragraph(topic_details["paragraph"])
    doc.save(file_name)


if __name__ == "__main__":
    try: 
        prompt = "Create a Word Document which showcases the company heirarchy. Comapny name is XYZ, it is an IT service company"
        docxDict = loads(generateDocx(prompt))
        print(docxDict)
        print(type(docxDict))
    except Exception as e:
        print("Hit an exception")
        print(e)
    generateDocxFromDict(docxDict)
